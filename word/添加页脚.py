# -*- coding:utf-8 -*-
#__author__:langzi
#__blog__:www.langzi.fun
import os

from docx import Document
document = Document(docx=os.path.join(os.getcwd(),'default.docx'))

section = document.sections[0]
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = "Report Power By Langzi"

document.add_paragraph('正文')
document.add_paragraph('我也是正文啊啊啊')

document.save('页脚.docx')